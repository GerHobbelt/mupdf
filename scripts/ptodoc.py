#! /usr/bin/env python3

'''
Takes XML output from 'mutool.py draw -F stext', and transforms into a .docx
file.

Args:

    -b <build-flags>
        <build-flags> is passed to ./scripts/mupdfwrap.py -b.

        Default is 'all' which builds mupdf.so and C++ wrappers and Python
        wrappers. The latter two are very slow, and usually not required after
        first run, in which case '-b m' will build just mupdf.so which usually
        builds very quickly.

    -d <directory>
        Set location of mupdf.so. Also passed to ./scripts/mupdfwrap.py -d.

        E.g. use .../build/shared-release for release build. Default is
        .../build/shared-debug.
Example:

    ./scripts/ptodoc.py -b m
    libreoffice ../ghostpdl/zlib/zlib.3.pdf.docx


Requirements:

    Packages:
         zip
         unzip
'''

import os
import subprocess
import sys
import xml.etree.ElementTree

import jlib
log = jlib.log

if 0:
    jlib.g_log_prefixes.append( jlib.LogPrefixFileLine())

mupdf_root = os.path.abspath( f'{__file__}/../../')
sys.path.append( os.path.abspath( f'{mupdf_root}/../python-docx/build/lib/'))
try:
    import docx
except ImportError as e:
    docx = None
    log( "warning: could not import python-docx's 'docx' module: {e}")


class PythonDocx:
    '''
    A .docx writer that uses python-docx's docx module.
    '''
    def __init__( self):
        self.document = docx.Document()
        self.paragraph = None
        self.run = None
        self.run_length = 0

    def new_paragraph( self):
        self.run = None
        self.run_length = 0
        self.paragraph = self.document.add_paragraph()

    def new_run( self, font_name, font_size_pt, bold=False, italic=False):
        self.run = self.paragraph.add_run()
        self.run.font.name = font_name
        self.run.font.size = docx.shared.Pt( font_size_pt)   # Convert to emu.
        self.run.font.bold = bold
        self.run.font.italic = italic
        self.run_length = 0

    def add_text( self, text):
        assert self.run
        self.run.add_text( text)
        self.run_length += len(text)

    def save( self, path):
        self.document.save( path)


class DirectDocx:
    '''
    A .docx writer that is standalone, except for using a template .docx
    document created by python-docx.
    '''
    def __init__( self):
        self.content = ''
        self.font_name = None
        self.font_size_pt = None
        self.run = None
        self.run_length = 0
        self.paragraph = None

    def new_paragraph( self):
        if self.run:
            # End previous run.
            self.content += '</w:t>'
            self.content += '</w:r>'
        if self.paragraph:
            # End previous paragraph.
            self.content += '\n</w:p>'
        self.font_name = None
        self.font_size_pt = None
        self.run = None
        self.run_length = 0
        self.content += '\n\n<w:p>'
        self.paragraph = True

    def new_run( self, font_name, font_size_pt, bold=False, italic=False):
        if (1
                and self.run
                and self.font_name == font_name
                and self.font_size_pt == font_size_pt
                and self.font_bold == bold
                and self.font_italic == italic
                ):
            return
        if self.run:
            self.content += '</w:t>'
            self.content += '</w:r>'
        self.content += '\n<w:r>'
        self.content += f'<w:rPr><w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        if bold:
            self.content += f'<w:b/>'
        if italic:
            self.content += f'<w:i/>'
        self.content += f'<w:sz w:val="{font_size_pt*2}"/>' # Looks like docx uses units of 0.5pt ?
        self.content += f'</w:rPr>'
        self.content += '<w:t xml:space="preserve">'
        self.font_name = font_name
        self.font_size_pt = font_size_pt
        self.font_bold = bold
        self.font_italic = italic
        self.run = True
        self.run_length = 0

    def add_text( self, text):
        assert self.run
        self.content += text
        self.run_length += len(text)

    def save( self, path, template=None, embed=True, preserve_dir=False):
        '''
        Modify a .docx file to contain our text.

        path:
            The docx file to modify.
        template:
            If None, we use python-docx to create a template .docx file.

            Otherwise the path of a .docx file to use as a template.
        embed:
            If true, we insert new text just after the '<w:body>' tag in the
            existing word/document.xml object in <path>.

            Otherwise we create a new word/document.xml object with our own XML
            header and footer.
        preserve_dir:
            If true, we leave <path>.dir/ containing the unzipped contents of
            the modified <path> file.
        '''
        tempdir = f'{path}.dir'
        jlib.system( f'rm -r {tempdir}', raise_errors=False, out=subprocess.DEVNULL)
        jlib.system( f'mkdir -p {tempdir}', prefix='    ')
        if template:
            jlib.system( f'unzip -q -d {tempdir} {template}', verbose=1, prefix='    ')
        else:
            docx.Document().save( path)
            jlib.system( f'unzip -q -d {tempdir} {path}', verbose=1, prefix='    ')

        if embed:
            with open( f'{tempdir}/word/document.xml') as f:
                embed_original = f.read()
            embed_marker = '<w:body>'
            embed_pos = embed_original.find( embed_marker)
            assert embed_pos >= 0
            embed_pos += len(embed_marker)

        # Overwrite word/document.xml with our content.
        #
        with open( f'{tempdir}/word/document.xml', 'w') as f:

            if embed:
                f.write( embed_original[ :embed_pos])

            else:
                # Write header.
                f.write( '<?xml version=\'1.0\' encoding=\'UTF-8\' standalone=\'yes\'?>\n')
                f.write( '<w:document')
                f.write( ' xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"')
                f.write( ' xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"')
                f.write( ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"')
                f.write( ' xmlns:mv="urn:schemas-microsoft-com:mac:vml"')
                f.write( ' xmlns:o="urn:schemas-microsoft-com:office:office"')
                f.write( ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"')
                f.write( ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"')
                f.write( ' xmlns:v="urn:schemas-microsoft-com:vml"')
                f.write( ' xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"')
                f.write( ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"')
                f.write( ' xmlns:w10="urn:schemas-microsoft-com:office:word"')
                f.write( ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')
                f.write( ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"')
                f.write( ' xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"')
                f.write( ' xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"')
                f.write( ' xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"')
                f.write( ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"')
                f.write( ' mc:Ignorable="w14 wp14"')
                f.write( '>\n')
                f.write( '<w:body>')

            # Write the content that we've built up so far.
            f.write( self.content)

            # Close any run and/or paragraph.
            if self.run:
                f.write( '</w:t>')
                f.write( '\n</w:r>')
            if self.paragraph:
                f.write( '\n</w:p>')
            f.write( '\n')

            if embed:
                f.write( embed_original[ embed_pos:])
            else:
                # Writer footer.
                f.write( '\n\n')
                f.write( '<w:sectPr w:rsidR="00FC693F" w:rsidRPr="0006063C" w:rsidSect="00034616">\n')
                f.write( '<w:pgSz w:w="12240" w:h="15840"/>\n')
                f.write( '<w:pgMar w:top="1440" w:right="1800" w:bottom="1440" w:left="1800" w:header="720" w:footer="720" w:gutter="0"/>\n')
                f.write( '<w:cols w:space="720"/>\n')
                f.write( '<w:docGrid w:linePitch="360"/>\n')
                f.write( '</w:sectPr>\n')
                f.write( '</w:body>\n')
                f.write( '</w:document>\n')

        jlib.system( f'cd {tempdir} && zip -q -r ../{os.path.basename(path)} .', verbose=1, prefix='    ')

        if not preserve_dir:
            jlib.system( f'rm -r {tempdir}', verbose=1, prefix='    ')



def walk_tree( root, *outputs):
    '''
    Walks tree from mupdf stext and applies various heuristics to join lines
    etc, generating information about text, runs, fonts etc which are sent to
    output objects.

    root:
        An xml.etree.ElementTree node.
    outputs:
        One or more objects with same methods as PythonDocx class.
    '''
    assert root.tag == 'document'

    for page in root:
        assert page.tag == 'page'

        for block in page:
            assert block.tag == 'block'

            # It's incorrect to always start new paragraph every time we see a
            # new block - e.g. paragraphs can span page boundaries. Prob need
            # to use a heuristic to look for paragraph start.
            #
            for output in outputs:
                output.new_paragraph()

            prev_line_ended_with_hyphen = False

            for line in block:
                assert line.tag == 'line'

                # If previous line did not end with hyphen, insert a space.
                #
                if not prev_line_ended_with_hyphen:
                    for output in outputs:
                        if output.run_length:
                            output.add_text( ' ')

                prev_line_ended_with_hyphen = False
            
                for font_i, font in enumerate(line):
                    assert font.tag == 'font'
                    line_end = (font_i + 1 == len(line))
                    font_name = font.get('name')
                    font_size_pt = float( font.get('size'))
                    for output in outputs:
                        output.new_run(
                                font_name,
                                font_size_pt,
                                '-Bold' in font_name,
                                '-Oblique' in font_name,
                                )
                
                    for char_i, char in enumerate(font):
                        assert char.tag == 'char'
                        font_end = (char_i + 1 == len(font))
                        c = char.get('c')
                    
                        if c == ' ' and char_i+1 < len(font):
                            # Discard spaces which overlap with the following
                            # character - these sometimes seem to appear in the
                            # middle of words.
                            #
                            def quad( c):
                                q = c.get('quad')
                                return [float(x) for x in q.split(' ')]
                            char_quad = quad( char)
                            ulx, uly, urx, ury, llx, lly, lrx, lry = char_quad

                            c2 = font[char_i+1]
                            char2_quad = quad( c2)
                            ulx2, uly2, urx2, ury2, llx2, lly2, lrx2, lry2 = char2_quad

                            omit = urx > ulx2
                            if omit:
                                pass
                                #log( '*** omitting space')
                            else:
                                pass
                                #log( 'found space')
                            if 0:
                                log( '    {c=} {c2=}')
                                log( '     {char_quad=}')
                                log( '    {char2_quad=}')
                                #log( '    {run_text=}')
                            if  omit:
                                # Ignore this ' ' character.
                                continue
                    
                        if line_end and font_end and c == '-':
                            # ignore '-' at end of line.
                            prev_line_ended_with_hyphen = True
                            continue
                        for output in outputs:
                            output.add_text( c)



def extract(extract_text_exe, mupdf_shared_dir, path_template, path_in, use_stext):
    '''
    Extracts text, and compares .docx's word/document.xml if reference file
    exists.
    '''
    log(f'Doing text extraction with {path_in}, use_stext={use_stext}')
    path_out = f'{path_in}-stext.docx' if use_stext else f'{path_in}.docx'
    path_content = f'{path_out}.content.xml'
    path_intermediate = f'{path_in}.intermediate.xml'

    # Run mutool.py to get intermediate xml.
    command = f'LD_LIBRARY_PATH={mupdf_shared_dir} PYTHONPATH={mupdf_shared_dir} scripts/mutool.py draw -F raw -o {path_intermediate} {path_in}'
    jlib.system( command, out=log, verbose=1, prefix='    ')

    command = (
                f' LD_LIBRARY_PATH=/home/jules/artifex/libbacktrace/.libs'
                f' MEMENTO_HIDE_MULTIPLE_REALLOCS=1'
                f' valgrind'
                f' ./{extract_text_exe}'
                f' -i {path_intermediate}'
                f' -t {path_template}'
                f' -p 1'    # preserve .docx temporary directory.
                f' -c {path_content}'
                f' -s {use_stext}'
                f' -o {path_out}'
                )
    jlib.system( command, out=log, verbose=1, prefix='    ')

    path_content = f'{path_out}.content.xml'
    path_content_ref = f'{path_out}.content.ref.xml'
    if os.path.exists(path_content_ref):
        jlib.system(f'diff -u {path_content_ref} {path_content}', out=log, verbose=1, prefix='    ')
    else:
        log(f'*** No reference content {path_content_ref} to compare with generated {path_content}')

    path_document_xml = f'{path_out}.dir/word/document.xml'
    path_document_xml_ref = f'{path_out}.word.document.ref.xml'
    if os.path.exists(path_document_xml_ref):
        jlib.system(f'diff -u {path_document_xml_ref} {path_document_xml}', out=log, verbose=1, prefix='    ')
    else:
        log(f'*** No reference document {path_document_xml_ref} to compare with generated {path_document_xml}')


def test(mupdf_shared_dir, so_build):

    if so_build:
        with jlib.LogPrefixScope('building mupdf.so: '):
            # Build mupdf.so and python wrapper.
            #
            command = f'./scripts/mupdfwrap.py -d build/shared-debug -b {so_build}'
            jlib.system( command, out=log, verbose=1, prefix='    ')
        
    
    with jlib.LogPrefixScope('building extract_text.exe: '):
        # Build extract_text.exe.
        #
        extract_text_c = 'source/tools/extract_text.c'
        extract_text_cc = 'source/tools/extract_text.c.c'
        extract_text_exe = 'extract_text.c.exe'

        memento_c = 'source/fitz/memento.c'
        memento_cc = 'source/fitz/memento.cc'
        if 0:
            jlib.build(
                    extract_text_c,
                    extract_text_cc,
                    f'cc -E -dD -g -o {memento_cc} -DMEMENTO {memento_c} -pthread -I include -I /usr/local/include -W -Wall -Wno-unused-parameter -Wno-unused-variable -Wno-unused-function build/shared-debug/libmupdf.so -Wl,--export-dynamic -L /usr/local/lib -lm -lexecinfo',
                    out=log,
                    )

            jlib.build(
                    extract_text_c,
                    extract_text_cc,
                    f'cc -E -dD -g -o {extract_text_cc} -DMEMENTO {extract_text_c} -pthread -I include -I /usr/local/include -W -Wall -Wno-unused-parameter -Wno-unused-variable -Wno-unused-function build/shared-debug/libmupdf.so -Wl,--export-dynamic -L /usr/local/lib -lm -lexecinfo',
                    out=log,
                    )

        command = (
                f'cc -g'
                f' -o {extract_text_exe}'
                #f' -DMEMENTO'
                f' {extract_text_c}'
                f' source/fitz/memento.c'
                f' -pthread'
                f' -I include'
                f' -I /usr/local/include'
                f' -W -Wall -Wno-unused-parameter -Wno-unused-variable -Wno-unused-function'
                f' build/shared-debug/libmupdf.so'
                f' -lm'
                )
        if os.uname()[0] == 'OpenBSD':
            command += ' -Wl,--export-dynamic -L /usr/local/lib -lexecinfo'
        else:
            command += ' -DHAVE_LIBDL -ldl'
        jlib.build(
                extract_text_c,
                extract_text_exe,
                command,
                out=log,
                )

    # Extract text from various input files.
    #
    path_template = '../Untitled1.docx'
    assert os.path.isfile(path_template), '*** We require an empty .docx document template called %s' % path_template
    for in_pdf in (
            f'{mupdf_root}/../ghostpdl/zlib/zlib.3.pdf',
            f'{mupdf_root}/../Python2.pdf',
            ):
        in_pdf_rel = os.path.relpath(in_pdf)
        for use_stext in 0, 1:
            with jlib.LogPrefixScope(f'{in_pdf_rel} use_stext={use_stext}: '):
                extract(
                        extract_text_exe,
                        mupdf_shared_dir,
                        path_template,
                        in_pdf_rel,
                        use_stext=use_stext,
                        )

    log( 'finished')


if __name__ == '__main__':
    args = iter(sys.argv[1:])
    so_build = 'all'
    mupdf_shared_dir = f'{mupdf_root}/build/shared-debug'
    while 1:
        try:
            arg = next(args)
        except StopIteration:
            break
        if arg == '-b':
            so_build = next(args)
        elif arg == '-d':
            mupdf_shared_dir = next(args)
        elif arg in ('-h', '--help'):
            print(__doc__)
        else:
            assert 0, 'unrecognised arg: %r' % arg
    try:
        test(mupdf_shared_dir, so_build)
    except Exception:
        print( jlib.exception_info())
        sys.exit(1)
